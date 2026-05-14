"""Word 문서 내보내기 서비스: 섹션별 초안 데이터를 .docx 파일로 변환한다."""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any


def _strip_inline_markup(text: str) -> str:
    """인라인 마크다운 기호(**bold**, *italic*, `code`)를 제거하고 순수 텍스트를 반환한다."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # **bold**
    text = re.sub(r"\*(.+?)\*", r"\1", text)       # *italic*
    text = re.sub(r"`(.+?)`", r"\1", text)          # `inline code`
    return text


def _add_content_to_doc(doc: Any, content: str) -> None:
    """마크다운 content 문자열을 줄 단위로 파싱해 Word 문서에 추가한다.

    처리 규칙:
    - ```...``` 코드 블록: 블록 전체 스킵 (기호 포함)
    - # / ## / ### 헤딩: Word Heading 2/3/4 스타일
    - - / * 글머리 기호: Word List Bullet 스타일
    - 1. / 1) 번호 목록: Word List Number 스타일
    - 빈 줄: 스킵 (연속 빈 줄 방지)
    - 일반 텍스트: 일반 단락, 인라인 기호 제거
    """
    in_code_block = False

    for line in content.splitlines():
        # 코드 블록 경계 감지
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue

        stripped = line.rstrip()

        # 헤딩: # / ## / ### / #### (들여쓰기 허용)
        m = re.match(r"^\s*(#{1,6})\s+(.*)", stripped)
        if m:
            # #→2, ##→3, ###→4, ####이상→4 (Word는 level 4까지)
            level = min(len(m.group(1)) + 1, 4)
            doc.add_heading(_strip_inline_markup(m.group(2).strip()), level=level)
            continue

        # 글머리 기호 목록: - 또는 * (들여쓰기 허용)
        m = re.match(r"^\s*[-*]\s+(.*)", stripped)
        if m:
            doc.add_paragraph(
                _strip_inline_markup(m.group(1).strip()),
                style="List Bullet",
            )
            continue

        # 번호 목록: 1. 또는 1) (들여쓰기 허용)
        m = re.match(r"^\s*\d+[.)]\s+(.*)", stripped)
        if m:
            doc.add_paragraph(
                _strip_inline_markup(m.group(1).strip()),
                style="List Number",
            )
            continue

        # 빈 줄 스킵
        if not stripped:
            continue

        # 일반 텍스트
        doc.add_paragraph(_strip_inline_markup(stripped))


def build_docx(project_name: str | None, sections: list[dict[str, Any]]) -> bytes:
    """섹션 목록을 받아 Word 문서 바이트를 반환한다.

    Args:
        project_name: 문서 제목으로 사용할 프로젝트명. None이면 기본 제목 사용.
        sections: [{"title": str, "content": str, "priority": int, ...}] 형태의 목록.

    Returns:
        .docx 파일 바이트.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title_text = project_name or "제안서 초안"
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for section in sections:
        section_title = section.get("title", "")
        content = section.get("content", "")

        doc.add_heading(section_title, level=1)
        _add_content_to_doc(doc, content)
        doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
